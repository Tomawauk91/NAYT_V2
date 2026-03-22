from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, color_hex):
    try:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._tc.get_or_add_tcPr().append(shading_elm)
    except:
        pass

document = Document()

# Title
title = document.add_heading('Rapport de Sécurité Complet', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph()

# Executive Summary Table
table = document.add_table(rows=4, cols=2)
table.style = 'Table Grid'

def style_row(row, bg_color):
    for cell in row.cells:
        set_cell_background(cell, bg_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

row = table.rows[0]
row.cells[0].text = 'Cible:'
row.cells[1].text = '{{ target }}'
style_row(row, 'E6F2FF')

row = table.rows[1]
row.cells[0].text = 'Date:'
row.cells[1].text = '{{ date }}'
style_row(row, 'E6F2FF')

row = table.rows[2]
row.cells[0].text = 'Client:'
row.cells[1].text = '{{ client_name }} ({{ client_company }})'
style_row(row, 'E6F2FF')

row = table.rows[3]
row.cells[0].text = 'Score CVSS Global:'
row.cells[1].text = '{{ overall_risk_score }} / 100'
style_row(row, 'E6F2FF')

document.add_paragraph()
document.add_heading('Résumé des Risques', level=1)

risk_table = document.add_table(rows=2, cols=5)
risk_table.style = 'Table Grid'

headers = ['Critique', 'Elevé', 'Moyen', 'Faible', 'Info']
counts = ['{{ critical_count }}', '{{ high_count }}', '{{ medium_count }}', '{{ low_count }}', '{{ info_count }}']
colors = ['FFCCCC', 'FFDAB9', 'FFFACD', 'E0FFFF', 'F0F8FF']

for i in range(5):
    cell = risk_table.cell(0, i)
    cell.text = headers[i]
    set_cell_background(cell, colors[i])
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cell_val = risk_table.cell(1, i)
    cell_val.text = counts[i]
    cell_val.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_background(cell_val, 'F9F9F9')

document.add_paragraph()
document.add_heading('Vulnérabilités Détaillées', level=1)
document.add_paragraph('{% if total_vulnerabilities > 0 %}')
document.add_paragraph('{% for vuln in vulnerabilities %}')

vuln_table = document.add_table(rows=4, cols=2)
vuln_table.style = 'Table Grid'

r0 = vuln_table.rows[0]
r0.cells[0].text = 'Nom'
r0.cells[1].text = '{{ vuln.title }}'
set_cell_background(r0.cells[0], 'F0F8FF')
r0.cells[0].paragraphs[0].runs[0].bold = True

r1 = vuln_table.rows[1]
r1.cells[0].text = 'Sévérité (CVSS)'
r1.cells[1].text = '{{ vuln.severity }} (Score: {{ vuln.cvss }})'
set_cell_background(r1.cells[0], 'F0F8FF')
r1.cells[0].paragraphs[0].runs[0].bold = True

r2 = vuln_table.rows[2]
r2.cells[0].text = 'Statut'
r2.cells[1].text = '{{ vuln.status }}'
set_cell_background(r2.cells[0], 'F0F8FF')
r2.cells[0].paragraphs[0].runs[0].bold = True

r3 = vuln_table.rows[3]
r3.cells[0].text = 'Description / Pre-Analysis'
r3.cells[1].text = '{{ vuln.description }}'
set_cell_background(r3.cells[0], 'F0F8FF')
r3.cells[0].paragraphs[0].runs[0].bold = True

document.add_paragraph()
document.add_paragraph('{% endfor %}')
document.add_paragraph('{% else %}')
document.add_paragraph('Aucune vulnérabilité majeure trouvée automatiquement. Cela ne garantit pas la sécurité totale. Effectuez des tests manuels.')
document.add_paragraph('{% endif %}')

document.save('/app/app/report_template.docx')
print('Enhanced template saved')
